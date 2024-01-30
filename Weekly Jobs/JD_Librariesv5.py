#Difference b/w JDLibraries1 being getDOB, extDOB and extageDOB  taken from PepHire Server 1
import pandas as pd
import re, subprocess, datetime
import shutil, docxpy
import win32com.client
import docx,PyPDF2
import slate3k as slate
from dateutil import parser
from difflib import SequenceMatcher
from blingfire import text_to_sentences
import cv2,os,docx2txt
from win32com.client import constants
import docx2python as docx2python
from docx2python.iterators import iter_paragraphs
from shutil import copyfile
from config import pephire_db, pephire_db_stat, Domain
from dateutil.relativedelta import relativedelta
#from calendar import month_abbr, month_name
#from datetime import datetime
#Define functions
#Function to find exact match of word in a sentence
from nltk.tokenize import WordPunctTokenizer
from nltk.collocations import BigramCollocationFinder
from nltk.metrics import BigramAssocMeasures

DegreeAbbs = pd.read_excel("DegreeAbbreviations.xlsx", sheet_name = "DegreeFormat")
converterpath = r"C:\Program Files\PDF2DOCX-CL\PDF2DOCX-CL.exe"

def get_bigrams(myString):
    tokenizer = WordPunctTokenizer()
    tokens = tokenizer.tokenize(myString)
    bigram_finder = BigramCollocationFinder.from_words(tokens)
    bigrams = bigram_finder.nbest(BigramAssocMeasures.chi_sq, 500)
    #Join the 2 tuples as a two word phrase
    bigrams = [' '.join(b) for b in bigrams]
    #Remove the leading and trialing spaces and ,
    bigrams = [b.strip(' ,') for b in bigrams]
    return pd.DataFrame(bigrams)

src_pdftotxt = r'C:\Pephire\pdftotext.exe'
def makeTable(table):
    data = []    
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    df = pd.DataFrame(data)
    return df
def fnString_exactSearch(string1, string2):
   if re.search(r"\b" + re.escape(string1) + r"\b", string2):
      return True
   return False
#Function to read word document in docx format
def getPDFTextasTxt_old(src_pdftotxt,FileName):
    txtFile = FileName
    txtFile = re.sub('.pdf','',txtFile)  
    Folderpath = os.path.dirname(FileName)
    dst = Folderpath + "\\pdftotext.exe"
    if(os.path.exists(dst)==False):
        copyfile(src_pdftotxt, dst) 
    subprocess.call([Folderpath + '\\pdftotext', FileName, txtFile+".txt"])
    text = open(txtFile+".txt",'r').read() 
    return text.replace('\n',' ').replace(',',' ')

#Function to read word document in docx format
def getPDFTextasTxt(FileName):
    destpath = re.sub('.pdf','.html',FileName)  
    subprocess.call(f'"{converterpath}" /src="{FileName}" /dst="{destpath}"', shell=True)
    docxpath = re.sub('.pdf','.docx',FileName)
    return getDocxText(docxpath)

def getDocxText(filename):
    return docxpy.process(filename)
def getDocxText_old(filename):
    try:        
        doc = docx2python.docx2python(filename)
        header_text = '\n'.join(iter_paragraphs(doc.header))
        footer_text = '\n'.join(iter_paragraphs(doc.footer))
    except:
        header_text = footer_text =""
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        try:
            fullText.append(para.text)
        except:
            pass
    fullText = '\n'.join(fullText)
    fullText = re.sub(r'[^\x00-\x7f]',r'', fullText)
    TableText = []
    tables = doc.tables
    tables = [makeTable(t) for t in tables]
    for tab in tables:
        for index, row in tab.iterrows(): 
            try:
                rowdata = list(set(row.values))
                rowdata = ' '.join(rowdata).strip()
                rowdata = re.sub(r'\s+',' ',rowdata)
                TableText.append(rowdata)
            except:
                pass

    TableText = '\n'.join(TableText)
    TableText = re.sub(r'[^\x00-\x7f]',r'', TableText)
    fullText = header_text + footer_text + fullText + ' ' + TableText
    return fullText.replace('\n',' ').replace(',',' ')

def getPDFText(filename):
    with open(filename,'rb') as f:
        extracted_text = slate.PDF(f)
    FullText = ""  
    extracted_text = pd.DataFrame(extracted_text)
    extracted_text.columns = ['Text']
    for i in range(extracted_text.shape[0]):
        sT = extracted_text['Text'][i]
        FullText = FullText+(sT)
    return re.sub(r'[^\x00-\x7f]',r'', FullText.replace('\n',' ').replace(',',' '))

def getPDFPyPDF(filename):
    pdfFileObj = open(filename, 'rb') 
    #print(name)
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
    NumPages = pdfReader.numPages
    txt = ""
    for i in range(0,NumPages):
        pageObj = pdfReader.getPage(i) 
        text = (pageObj.extractText()) 
        txt = txt+" "+text
    txt = ''.join([i if ord(i) < 128 else '' for i in txt])
    txt = txt.replace('\n','')
    pdfFileObj.close()
    return txt.replace('\n',' ').replace(',',' ')
def getPDF(filename):
    text = ""
    try:        
        text = getPDFPyPDF(filename)
    except:
        try:
            text = getPDFText(filename)
        except:
            text = ""
    if len(text)<10:
        text = getPDFPyPDF(filename)    
    return text.replace('\n',' ').replace(',',' ')


# format date text and return date only
def extDOB(Line):
    # remove unnecessary chars
    sText = Line.strip(" :\t-")
    sText = re.sub(r'[\-/\.\,]', ' ', sText)
    sText = re.sub(r'(\d+)(th|nd|rd|st)', r'\1 ', sText)
    sText = sText.replace('–', ' ')
    sText = re.sub(r'\s+', ' ', sText).strip(' .')
    sText = sText.replace(' th ',' ')
    # check for any text patterns
    dMat = re.search(r'\d+\s\d+\s\d+|\d+\s[a-zA-Z]+\s\d+', sText)
    if dMat:
        try:
            # parse date
            dateText = parser.parse(str(dMat.group()))
            return str(dateText).replace(' 00:00:00','')
        except:
            return ''
    return '' 

# Get DOB from the string
def getDOB(Line):    
    detect = False
    sText = ""
    # List of possible DOB patterns
    patterns = ["date of birth", "dob ", "dob:", 'd.o.b', 'born', "birth date"]
    for p in patterns:
        if p in Line:
            sText = Line.replace(p,"")
            detect = True
            break
    # format date text and return date only 
    dateText = extDOB(sText)
    return dateText, detect

# extract dob if age is given
def extageDOB(Line):
    #print(Line)
    # check if age is in the string
    if ' age ' in Line.lower() or Line.lower().startswith('age') or Line.lower().endswith(' age'):
        # extract numbers
        sText = re.findall(r'[0-9]{2}',Line)
        # convert into date
        sText = datetime.now() - relativedelta(years = int(sText[0]))
        sText = sText.strftime('%Y-%m-%d')
        # return with a Flag
        return sText, True
    else:
        return '', False

def getPassport(Line):
    sText, Flag = "", False
    if(("passport" in Line)):
        sText = re.findall(r'[a-z0-9]+[\s\t]*?$',Line)
        if len(sText)>0:
            sText = re.findall(r'[a-z0-9]+',sText[0])[0]
        Flag = True
    '''
    if(("passport number:" in Line)):
        #sText = Line.replace("passport number:","").strip(" :")#passport number:
        sText = re.sub(r'(passport\snumber)(.*)', r'\2 ', Line)
        sText = re.sub(r'([^a-zA-Z0-9]+)(.*)', r'\2 ', sText)
        Flag = True
    elif(("passport number" in Line)):
        sText = re.sub(r'(passport\snumber)(.*)', r'\2 ', Line)
        sText = re.sub(r'([^a-zA-Z0-9]+)(.*)', r'\2 ', sText)
        Flag = True
    elif(("passport num" in Line)):
        sText = re.sub(r'(passport\snum)(.*)', r'\2 ', Line)
        sText = re.sub(r'([^a-zA-Z0-9]+)(.*)', r'\2 ', sText)
        Flag = True
    elif(("passport no" in Line)):
        sText = re.sub(r'(passport\sno)(.*)', r'\2 ', Line)
        sText = re.sub(r'([^a-zA-Z0-9]+)(.*)', r'\2 ', sText)
        Flag = True
    elif(("passport #" in Line)):
        sText = re.sub(r'(passport #)(.*)', r'\2 ', Line)
        sText = re.sub(r'([^a-zA-Z0-9]+)(.*)', r'\2 ', sText)
        Flag = True
    elif(("passport#" in Line)):
        sText = re.sub(r'(passport#)(.*)', r'\2 ', Line)
        sText = re.sub(r'([^a-zA-Z0-9]+)(.*)', r'\2 ', sText)
        Flag = True
    elif(("passport" in Line)):
        sText = re.sub(r'(passport)(.*)', r'\2 ', Line)
        sText = re.sub(r'([^a-zA-Z0-9]+)(.*)', r'\2 ', sText)
        Flag = True
    elif(("passport:" in Line)):
        sText = re.sub(r'(passport:)(.*)', r'\2 ', Line)
        sText = re.sub(r'([^a-zA-Z0-9]+)(.*)', r'\2 ', sText)
        Flag = True
    '''
    #Get the string matching passport number pattern
    if(sText!=""):
        try:
            sText = re.findall(r"[a-z0-9]+", sText)
            sText = sText[0]
        except:
            #print("Passport exception")
            sText=""
    return sText, Flag

def extPassport(sText):
    sText = re.findall(r"[a-z0-9]+", sText)
    return sText[0]

def getMaritalStatus(Line):
    sText, Flag = "", False
    '''
    if "marital\sstat" in Line:
        sText = re.sub(r'marital\sstat[^\s]+', '',Line)
        #sText = Line.replace("marital status","")
        sText = sText.strip(" :\t-")
        Flag = True
    '''
    if "relationship" or "marital" in Line:
        sText = re.findall(r'[a-zA-Z]+[\s\t]*$',Line)
        if len(sText)>0:
            sText = re.findall(r'[a-z]+',sText[0])[0]
        #sText = Line.replace("marital status","")
        #sText = sText.strip(" :\t-")
        Flag = True
    if(sText.lower() == "married" or sText.lower()=="single"):
        return sText, Flag
    else:
        return "", Flag

def extMaritalStatus(sText):
    sText = sText.strip(" :\t-")
    if(sText.lower() == "married" or sText.lower()=="single"):
        return sText
    else:
        return ""
#Need to find the biggest number from the list of values that contains year of exp. Remove 'year' and '+' to get the numeric value of year
def getMaxYear(t1):   
    if(len(t1)==0):
        return 0
    Yr = pd.DataFrame(t1).reset_index()
    Yr.columns = ['Index','Text']
    
    for i in range(Yr.shape[0]): 
        Yr['Text'][i] = Yr['Text'][i].replace('year','').replace('+','')
        
    Yr.Text = Yr.Text.astype(float)
    Yr = Yr[Yr.Text < 50]
    maxYear = Yr['Text'].max()
    return maxYear
'''
def getVisaStatus(Line):
    sText = ""
    VisaFlag = False
    LineL = Line.lower()
    if(("visa" in LineL)):
        #sText = Line.replace("visa","").strip(" :|-")
        sText = re.sub(r'(visa)(.*)', r'\2 ', LineL)
        sText = re.sub(r'([^a-zA-Z]+)(.*)', r'\2 ', sText)
        if(len(sText)>50):
            sText =""
        VisaFlag = True
    if(("visas" in LineL)):
        #sText = Line.replace("visa","").strip(" :|-")
        VisaFlag = True
        sText = re.sub(r'(visas)(.*)', r'\2 ', LineL)
        sText = re.sub(r'([^a-zA-Z]+)(.*)', r'\2 ', sText)
        if(len(sText)>50):
            sText =""
        VisaFlag = True
    if(("work\spermit" in LineL)):
        #sText = Line.replace("visa","").strip(" :|-")
        VisaFlag = True
        sText = re.sub(r'(work\spermit)(.*)', r'\2 ', LineL)
        sText = re.sub(r'([^a-zA-Z]+)(.*)', r'\2 ', sText)
        if(len(sText)>50):
            sText =""
        VisaFlag = True
    return sText.title().strip(" :\t-"), VisaFlag
'''
def getVisaStatus(Line):
    sText = ""
    VisaFlag = False
    LineL = Line.lower()
    if(("visa" in LineL)):
        #sText = Line.replace("visa","").strip(" :|-")
        try:
            sText = re.findall(r'(visas?[^a-z0-9]+)(.*$)',LineL)[0][1].strip(' \t')
        except:
            pass
        sText = re.sub(r'visa','',sText).strip(' \t')
        if(len(sText)>50):
            sText =""
        VisaFlag = True
    if(("visas" in LineL)):
        #sText = Line.replace("visa","").strip(" :|-")
        try:
            sText = re.findall(r'(visas[^a-z0-9]+)(.*$)',LineL)[0][1].strip(' \t')
        except:
            pass
        #sText = re.sub(r'visas','',sText).strip(' \t')        
        if(len(sText)>50):
            sText =""
        VisaFlag = True
    if(("work\spermit" in LineL)):
        #sText = Line.replace("visa","").strip(" :|-")
        try:
            sText = re.findall(r'(work\spermit[^a-z0-9]+)(.*$)',LineL)[0][1].strip(' \t')
        except:
            pass
        #sText = re.sub(r'work\spermit','',sText).strip(' \t')
        if(len(sText)>50):
            sText =""
        VisaFlag = True
    return sText.title().strip(" :\t-"), VisaFlag

def extVisaStatus(Line):
    sText = ""
    sText = re.findall(r"[a-zA-Z]+.*", Line)[0]
    if len(sText) < 50 and sText != "":
        return sText.title().strip(" :\t-")
    else:
        return ""
def similar(a, b):
    return SequenceMatcher(None, a, b).ratio()


def getSex(text):
    text = text.lower()
    sents = text_to_sentences(text).split('\n')
    for s in sents:
        if 'female' in s:
            return 'Female'
        elif 'male' in s:
            return 'Male'
        elif 'gender' in s:
            if 'f' in s:
                return 'Female'
            else:
                return 'Male'
        elif 'sex' in s:
            if 'f' in s:
                return 'Female'
            else:
                return 'Male'
            
    text = text.title()  
    for s in sents:
        if 'Mrs.' in s or 'Ms.' in s:
            return 'Female'
        elif 'Mr.' in s:
            return 'Male'



def FaceRecognition(imagePath):    
    cascPath = "haarcascade_frontalface_default.xml"
    faceCascade = cv2.CascadeClassifier(cascPath)
    image = cv2.imread(imagePath)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    faces = faceCascade.detectMultiScale(
        gray,
        scaleFactor=1.1,
        minNeighbors=5,
        minSize=(30, 30),
        flags = cv2.CASCADE_SCALE_IMAGE
    )
    if len(faces)>0:
        return 1
    else:
        return 0
    print("Found {0} faces!".format(len(faces)))
    for (x, y, w, h) in faces:
        cv2.rectangle(image, (x, y), (x+w, y+h), (0, 255, 0), 2)    
    cv2.imshow("Faces found", image)
    cv2.waitKey(0)
    
def extract_jpg_from_PDF(file,outputpath):
    try:
        os.mkdir(outputpath)
    except:
        pass
    pdf = open(file,'rb')
    pdf = pdf.read()
    startmark = b"\xff\xd8"
    startfix = 0
    endmark = b"\xff\xd9"
    endfix = 2
    i = 0
    writename = re.search(r'([^\\]+).pdf$',file).group(1)
    njpg = 1
    writelist = []
    while True:
        istream = pdf.find(b"stream", i)
        if istream < 0:
            #print("Y")
            break
        istart = pdf.find(startmark, istream, istream + 20)
        if istart < 0:
            #print("A")
            i = istream + 20
            continue
        iend = pdf.find(b"endstream", istart)
        if iend < 0:
            raise Exception("Didn't find end of stream!")
        iend = pdf.find(endmark, iend - 20)
        if iend < 0:
            raise Exception("Didn't find end of JPG!")

        istart += startfix
        iend += endfix
        #print("JPG %d from %d to %d" % (njpg, istart, iend))
        jpg = pdf[istart:iend]
        
        with open(f"{outputpath}\{writename}%d.jpg" % njpg, "wb") as jpgfile:
            jpgfile.write(jpg)
            writelist.append(jpgfile)
        if FaceRecognition(f"{outputpath}\{writename}%d.jpg" % njpg)!=0:
            return f"{outputpath}\{writename}%d.jpg" % njpg
        njpg += 1
        i = iend
    return "No Images"

def extract_jpg_from_docx(file,outputpath):
    writename = re.search(r'([^\\]+).docx$',file).group(1)
    # extract text
    #text = docx2txt.process(file)
    path = f'{outputpath}\{writename}'
    try:
        shutil.rmtree(path, ignore_errors=False, onerror=None)
    except:
        pass
    try:
        os.makedirs(path)
    except:
        pass
    # extract text and write images in /tmp/img_dir
    docx2txt.process(file, path) 
    outfiles = [filenames for (dirpath, dirnames, filenames) in os.walk(path)][0]
    writelist = []
    for i in outfiles:
        name = re.sub(r'image',writename,i)
        try:
            os.rename(f'{outputpath}\{writename}\{i}',f'{outputpath}\{name}')
        except:
            pass
        writelist.append(f'{outputpath}\{name}')
    os.removedirs(f'{outputpath}\{writename}')
    for i in writelist:
        if (FaceRecognition(i)!= 0):
            return i
    return "No Images"
def save_as_docx(path):
    # Opening MS Word
    word = win32com.client.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(path)
    doc.Activate ()

    # Rename path with .docx
    new_file_abs = os.path.abspath(path)
    new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)

    # Save and Close
    word.ActiveDocument.SaveAs(
        new_file_abs, FileFormat=constants.wdFormatXMLDocument
    )
    doc.Close(False)

def getImage(file,outputpath):
    try:
        fileName = re.search(r'([a-zA-Z0-9]+)(\.[a-z]+)', file).group(1)
        result = [os.path.join(dp, f) for dp, dn, filenames in os.walk(outputpath) for f in filenames]
        resultFiles = [f for f in result if fileName in f]
        
        for f in resultFiles:
            os.remove(f)
    except:
        pass
    
    txt=""
    try:
        if file.endswith(".pdf"):
            txt = extract_jpg_from_PDF(file,outputpath)
        elif file.lower().endswith(".docx"):
            if not file.endswith(".docx"):
                dst = re.sub(".{3}$",".docx",file) 
                os.rename(file, dst)
            txt = extract_jpg_from_docx(file,outputpath)
        elif file.lower().endswith(".doc"):
            save_as_docx(file)
            os.remove(file)
            file1 = re.sub(r'.doc',r'.docx',file)
            extract_jpg_from_docx(file1,outputpath)
          
    except:
        txt=""
    return txt

def CheckJob(txt):
    TextLower = txt.lower()
    JobsMatched = [i for i in emplist if i in TextLower]
    TextExtracted = []
    for i in JobsMatched:
        #print(i)
        w = re.findall(r'[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s?'+re.escape(i)+'\s?[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+',TextLower)  
        #print(w)
        TextExtracted = TextExtracted + w
    for i in JobsMatched:
        for j,k in enumerate(TextExtracted):
            TextExtracted[j] = re.sub(r'\s?'+re.escape(i)+'\s?',' ',TextExtracted[j])
            
    for i,j in enumerate(TextExtracted):
        dates = search_dates(TextExtracted[i])
        try:
            dates = [k[0] for k in dates]
            #print(dates)
            for r in dates:
                if len(r) <= 2:
                    r = f" {r} "
                TextExtracted[i] = re.sub(r,'',TextExtracted[i])  
        except:
            pass
        TextExtracted[i] = re.sub(r'\s+',' ',TextExtracted[i])
        words = word_tokenize(TextExtracted[i])
        words = [w for w in words if not w in stops]
        TextExtracted[i] = ' '.join(words)
        x = word_tokenize(TextExtracted[i])
        x1 = nltk.pos_tag(x)
        TextExtracted[i] = ' '.join([k[0] for k in x1 if k[1]=="NNS" or k[1]=="NN" ])
    return TextExtracted
def CheckCap(com):
    if len(com) <= 2:
        return False
    elif len(com) < 4:
        if not com.isupper():
            return False
    return True        
def substringSieve(string_list):
    string_list.sort(key=lambda s: len(s), reverse=True)
    out = []
    for s in string_list:
        if not any([s in o for o in out]):
            out.append(s)
    return out
                
# importing required modules 
def ForPdf(name):
    try:
        pdfFileObj = open(name, 'rb') 
        #print(name)
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
        NumPages = pdfReader.numPages
        txt = ""
        for i in range(0,NumPages):
            pageObj = pdfReader.getPage(i) 
            text = (pageObj.extractText()) 
            txt = txt+" "+text
        txt = ''.join([i if ord(i) < 128 else '' for i in txt])
        txt = txt.replace('\n','')
        pdfFileObj.close()
        txt = re.sub(r'\s+',' ',txt)
        txt = re.sub(r'[^\x00-\x7F]+',' ', txt)
        txt = re.sub(r'([a-zA-Z]+)\-([a-zA-Z]+)',r'\1\2',txt)
        TextExtracted = CheckJob(txt)
        txt = re.sub(r'Software|Solutions|Tester|Developer|Create','',txt)
        txt = re.sub(r'[^\x00-\x7F]+',' ', txt)
        txt = re.sub(r'\s+',' ',txt)
        #txt = re.sub(r'[^a-z^A-Z^0-9^\.^\&^\:^\"^\'^\s]','',txt)
        #txt = txt.encode('windows-1252').decode().strip()
        txt = re.sub(r'([a-zA-Z]+)\-([a-zA-Z]+)',r'\1\2',txt)
        
        doc = nlp(txt)
        orgList = []
        for entity in doc.ents:
            if entity.text != " ":
                if entity.label_ == "ORG" or "PERSON":
                    orgList.append(entity.text)
        w1 = re.findall(r'[21][0-9]{3}',txt.lower())
        w2 = re.findall(r'year\'?s?\'?|yr\'?s?\'?',txt.lower())
        w1 = w2 + w1
        w1 = list(set(w1))
        w2 = []
        w3 = []
        w4 = []
        w5 = []
        w6 = []
        w7 = []
        w8 = []
        w3 = w3 + re.findall(r'\sas\s[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\sin\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+' ,txt)
        w4 = w4 + re.findall(r'\sin\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?as\s' ,txt)
        w5 = w5 + re.findall(r'\sas\s[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\sat\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+' ,txt)
        w6 = w6 + re.findall(r'\sat\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?as\s' ,txt)
        w7 = w5 + re.findall(r'\sas\s[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\swith\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+' ,txt)
        w8 = w6 + re.findall(r'\swith\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?as\s' ,txt)
        for i,j in enumerate(w1):
            w2 = w2 + re.findall(r'\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s?'+re.escape(j)+r'\s?[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+' ,txt)
            
        w3 = w3 + w2 + w4 + w5 + w6 + w7 + w8
        txt2 = ' '.join(w3) 
        orgListVerified = []
        for a in orgList:
            if a in txt2:
                orgListVerified.append(a)
        orgListVerified = list(set(orgListVerified))
        orgListFinal = []
        for i,j in enumerate(orgListVerified):
            for k,l in enumerate(TextExtracted):
                if j.lower() in l:
                    orgListFinal.append(j)
        orgListFinal = list(set(orgListFinal))
        orgListFinal = substringSieve(orgListFinal)
    except:
        pass
    
def Fordocx(name):
    try:
        txt = getDocxText(name)
        txt = re.sub(r'\s+',' ',txt)
        txt = re.sub(r'[^\x00-\x7F]+',' ', txt)
        txt = re.sub(r'([a-zA-Z]+)\-([a-zA-Z]+)',r'\1\2',txt)
        TextExtracted = CheckJob(txt)
        txt = re.sub(r'Software|Solutions|Tester|Developer|Create|Hackathon|HACKATHON','',txt)
        txt = re.sub(r'[^\x00-\x7F]+',' ', txt)
        txt = re.sub(r'\s+',' ',txt)
        #txt = re.sub(r'[^a-z^A-Z^0-9^\.^\&^\:^\"^\'^\s]','',txt)
        #txt = txt.encode('windows-1252').decode().strip()
        txt = re.sub(r'([a-zA-Z]+)\-([a-zA-Z]+)',r'\1\2',txt)
        TextExtracted = CheckJob(txt)
        doc = nlp(txt)
        orgList = []
        for entity in doc.ents:
            if entity.text != " ":
                if entity.label_ == "ORG" or "PERSON":
                    orgList.append(entity.text)
        orgList = list(set(orgList))

        w1 = re.findall(r'[21][0-9]{3}',txt.lower())
        w2 = re.findall(r'year\'?s?\'?|yr\'?s?\'?',txt.lower())
        w1 = w2 + w1
        w1 = list(set(w1))
        w2 = []
        w3 = []
        w4 = []
        w5 = []
        w6 = []
        w7 = []
        w8 = []
        w3 = w3 + re.findall(r'\sas\s[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\sin\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+' ,txt)
        w4 = w4 + re.findall(r'\sin\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?as\s' ,txt)
        w5 = w5 + re.findall(r'\sas\s[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\sat\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+' ,txt)
        w6 = w6 + re.findall(r'\sat\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?as\s' ,txt)
        w7 = w5 + re.findall(r'\sas\s[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\swith\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+' ,txt)
        w8 = w6 + re.findall(r'\swith\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?[^\s]+?\s?as\s' ,txt)
        for i,j in enumerate(w1):
            w2 = w2 + re.findall(r'\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s?'+re.escape(j)+r'\s?[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+\s[^\s]+' ,txt)
            
        w3 = w3 + w2 + w4 + w5 + w6 + w7 + w8
        txt2 = ' '.join(w3) 
        doc = nlp(txt2)
        orgList1 = []
        for entity in doc.ents:
            if entity.text != " ":
                if entity.label_ == "ORG" or "PERSON":
                    orgList1.append(entity.text)
        orgList1 = list(set(orgList1))
        orgListVerified = []
        for a in orgList:
            if a in txt2:
                orgListVerified.append(a)
        orgListVerified = list(set(orgListVerified))
        orgListFinal = []
        for i,j in enumerate(orgListVerified):
            for k,l in enumerate(TextExtracted):
                if j.lower() in l:
                    orgListFinal.append(j)
        orgListFinal = list(set(orgListFinal))
        orgListFinal = substringSieve(orgListFinal)
    except:
        pass
def getCompany(name):
    if name.endswith(".pdf"):
        orgFinal = ForPdf(name)
    elif name.endswith(".docx"):
        orgFinal = Fordocx(name)
    return orgFinal

#Text Modifications
def TextModifications(txt):
    txt = txt.replace('\n',' ')                             #Replacing \n with \s
    txt = re.sub(r'\s+',' ',txt)                            #Replacing multiple spaces with single space
    txt = re.sub(r'[^\x00-\x7F]+',' ', txt)                 #Replacing multiple spaces with single space for unicode
    txt = re.sub(r'([a-zA-Z]+)\-([a-zA-Z]+)',r'\1\2',txt)   #Word1-Word2 to Word1Word2
    return txt

def isNaN(num):
    try:
        a = num
        return 0
    except:
        return 1
def getEducationalQualification(txt):
    name = 'name'
    txt = txt.replace('’',"'")
    WordsList = pd.DataFrame(re.split('\s|\(|\)|\-|:|\[|\]',txt), columns = ["Original"])
    WordsList['Formatted'] = WordsList.Original.str.replace('.', '')
    WordsList['Formatted'] = WordsList.Formatted.str.replace(',', '')
    WordsList['Formatted'] = WordsList['Formatted'].str.lower()
    Out = pd.merge(WordsList, DegreeAbbs, left_on = "Formatted", right_on = "DegreeLower", how = 'outer', indicator = True )
    Out1 = Out[Out['_merge']=='both']
    Out1['name'] = name
    Out1 = Out1[['name','Original','Degree']]
    WordsList['Original'] = WordsList.Original.str.replace('.', '')
    WordsList['Original'] = WordsList.Original.str.replace(',', '')
    for i,j in enumerate(WordsList['Original']):
        if j=='B' or j=='M' :
            temp = WordsList['Original'][i] + WordsList['Original'][i+1]
            tempout = WordsList['Formatted'][i] + WordsList['Formatted'][i+1]
            if tempout.lower() in list(DegreeAbbs['DegreeLower']):
                Out1 = Out1.append(pd.Series([name, temp], index=['name','Original']), ignore_index=True)
    for i,j in enumerate(WordsList['Original']):
        if j=='Bachelor' or j=='Bachelors' or j == 'Master' or j == 'Masters' or j == "Bachelor's" or j == "Master's" or j.lower() == 'post':
            temp = WordsList['Original'][i] + ' ' + WordsList['Original'][i+1] + ' ' + WordsList['Original'][i+2] + ' ' + WordsList['Original'][i+3]
            tempout = WordsList['Formatted'][i] + ' ' + WordsList['Formatted'][i+1] + ' ' + WordsList['Formatted'][i+2] + ' ' + WordsList['Formatted'][i+3]
            Out1 = Out1.append(pd.Series([name, temp], index=['name','Original']), ignore_index=True)
    for i in DegreeAbbs['FullForm']:
        if i in txt:
            #print({'name':[name] , 'Original':[i]})
            Out1 = Out1.append(pd.Series([name, i], index=['name','Original']), ignore_index=True)
    Out1 = Out1[Out1.Original.str.islower()==0]
    temp = []
    Out1 = Out1.reset_index().drop('index',axis = 1)
    for i,j in Out1.iterrows():
        if type(Out1['Degree'][i]) != float and Out1['Original'][i] != Out1['Degree'][i] and not Out1['Original'][i].isupper():
            temp.append(i)
    #Delete wrong entries
    Out1 = Out1.drop(temp)
    Out1['Three'] = Out1['Original']
    for i,row in Out1.iterrows():
        templist = str(Out1['Three'][i]).split()
        if len(templist) > 3:
            templist = templist[0:3]
            Out1['Three'][i] = ' '.join(templist)
    Out1.Degree.fillna(Out1.Original, inplace=True)
    Out1['Education'] = Out1['Degree']
    for i, row1 in Out1.iterrows():
        flag = 0
        templist = str(Out1['Education'][i]).split()
        if len(templist) >= 3:
            for j, row2 in DegreeAbbs.iterrows():
                if similar(Out1['Three'][i],DegreeAbbs['FullForm'][j]) > 0.85:
                    #print(name, ":", Out1['Three'][i], ":", DegreeAbbs['Degree'][j],":",DegreeAbbs['FullForm'][j],":",similar(Out1['Three'][i],DegreeAbbs['FullForm'][j]))
                    flag = 1
                    Out1['Education'][i] = DegreeAbbs['Degree'][j]
                    break
            if flag==0:
                for j, row2 in DegreeAbbs.iterrows():
                    if similar(Out1['Education'][i],DegreeAbbs['FullForm'][j]) > 0.85:
                        #print(name, ":", Out1['Three'][i], ":", DegreeAbbs['Degree'][j],":",DegreeAbbs['FullForm'][j],":",similar(Out1['Three'][i],DegreeAbbs['FullForm'][j]))
                        flag = 1
                        Out1['Education'][i] = DegreeAbbs['Degree'][j]
                        break
                if flag == 0:
                    TwoWord = ' '.join(templist[0:2])
                    for j, row2 in DegreeAbbs.iterrows():
                        if similar(TwoWord.lower(),DegreeAbbs['FullForm'][j].lower()) > 0.85:
                            #print(TwoWord, ' ', DegreeAbbs['FullForm'][j])
                            #print(name, ":", Out1['Three'][i], ":", DegreeAbbs['Degree'][j],":",DegreeAbbs['FullForm'][j],":",similar(Out1['Three'][i],DegreeAbbs['FullForm'][j]))
                            flag = 1
                            Out1['Education'][i] = DegreeAbbs['Degree'][j]
                            break
            
    Out1 = Out1[['name','Education']]
    Out1 = Out1[Out1['Education'].map(len) < 16]
    Out1 = Out1.drop_duplicates() 
    tempM = []
    MasterPosition = []
    if 'Master' in list(Out1.Education) or 'Masters' in list(Out1.Education) or "Master's" in list(Out1.Education):        
        MasterList = ['Master','Masters',"Master's"]
        for i, row in Out1.iterrows():
            if Out1['Education'][i].startswith('M') and Out1['Education'][i] not in MasterList:
                tempM.append(i)
            if Out1['Education'][i] in MasterList:
                MasterPosition.append(i)
    if len(tempM)>0 and len(MasterPosition)>0:
        Out1 = Out1.drop(MasterPosition)
    tempB = []
    BachelorPosition = []
    if 'Bachelor' in list(Out1.Education) or 'Bachelors' in list(Out1.Education) or "Bachelor's" in list(Out1.Education):        
        BachelorList = ['Bachelor','Bachelors',"Bachelor's"]
        for i, row in Out1.iterrows():
            if Out1['Education'][i].startswith('B') and Out1['Education'][i] not in BachelorList:
                tempB.append(i)
            if Out1['Education'][i] in BachelorList:
                BachelorPosition.append(i)  
    if len(tempB)>0 and len(BachelorPosition)>0:
        Out1 = Out1.drop(BachelorPosition) 
    Out = list(Out1.Education)[0:3]
    Out = ','.join(Out) 
        
    return(Out)
    
def get_JDExp(jd):
    jd = str(jd).lower()
    exp = re.search(r'(experience|exp)[\s:]+(\d+)[\-\sto]+(\d+)\s(yrs|yr|years)', jd)
    if exp:
        return round((int(exp.group(2)) + int(exp.group(3)))/2, 2)
    exp = re.search(r'(experience|exp)[\s:]+(\d+|\d+\.\d+)\s?(yrs|yr|years)', jd)
    if exp:
        return float(exp.group(2))
    
    exp = re.search(r'(\d+)[\-\sto]+(\d+)\s(yrs|yr|years|year)[ofworking\s]+exp', jd)
    if exp:
        return round((int(exp.group(1)) + int(exp.group(2)))/2, 2)
    exp = re.search(r'(\d+|\d+\.\d+)\+?\s?(yrs|yr|years|year)[ofworking\s]+exp', jd)
    if exp:
        return float(exp.group(1))
    
    lines = jd.splitlines()
    lines = [l.strip() for l in lines if l.strip()]
    
    for i in range(len(lines)):
        if re.search(r'(\d+)[\-\sto]+(\d+)', lines[i]) and 'exp' in lines[i]:
            exp = re.search(r'(\d+)[\-\sto]+(\d+)', lines[i])
            return round((int(exp.group(1)) + int(exp.group(2)))/2, 2)
        elif re.search(r'(\d+|\d+\.\d+)\+?\s?(yrs|yr|years|year)', lines[i]) and 'exp' in lines[i]:
            exp = re.search(r'\d+|\d+\.\d+', lines[i])
            return float(exp.group())
        elif lines[i]=='experience':
            exp = re.search(r'(\d+)[\-\sto]+(\d+)\s(yrs|yr|years|year)', lines[i+1])
            if exp:
                return round((int(exp.group(1)) + int(exp.group(2)))/2, 2)
            exp = re.search(r'(\d+|\d+\.\d+)\+?\s?(yrs|yr|years|year)', lines[i+1])
            if exp:
                return float(exp.group(1))
    return 0

# -*- coding: utf-8 -*-
"""
Profile sensing code to get skill and top3 skills
"""
import pandas as pd
import mysql.connector,re
from config import pephire_db_stat, Domain
#from JD_Libraries import get_bigrams
#Get bigrams

#load the skillwords that are clustered. This is saved in DB based on daily or weekly jobs that pull data from naukri
mydb_static = mysql.connector.connect(**pephire_db_stat)
mycursor_static = mydb_static.cursor()
if Domain == 1:
    df_MainCluster = pd.read_sql("SELECT * FROM pephire_static.maincluster;", mydb_static)
    df_SkillNorm = pd.read_sql("SELECT * FROM pephire_static.newskillnorm;", mydb_static)    
else:
    #SkillMap = pd.read_excel('UST - SkillMapping.xlsx',sheet_name = 'Mapper')
    df_MainCluster = pd.read_sql("SELECT * FROM pephire_static.maincluster_retail_v1;", mydb_static)
    df_SkillNorm = pd.read_sql("SELECT * FROM pephire_static.newskillnorm_retail_v1;", mydb_static)
try:
    df_MainCluster = df_MainCluster.drop(["row_names"], axis = 1)
    df_SkillNorm = df_SkillNorm.drop(["row_names"], axis = 1)
except:
    pass

df_SkillMap = pd.read_sql("SELECT * FROM pephire_static.skillmap;", mydb_static)
mydb_static.close()

df_MainClusterUniqueSkills = df_MainCluster[['Skill']].drop_duplicates().reset_index()
df_MainClusterUniqueSkills = df_MainClusterUniqueSkills[['Skill']]
df_MainClusterUniqueSkills.columns = ['Word']

def GetSkill_New(CurrentFile,sDocText):    
    try:
        ResumeWords = pd.DataFrame(re.split(' ',sDocText.lower().replace('\n',' ').replace('.',' ').replace(',',' ').replace('/',' ').replace(':',' ').replace('-',' ').replace('?',' ').replace('_',' ')))
        ResumeWords.columns = ['Word']
        ResumeWords['Location'] = ResumeWords.index
        ResumeWords.columns =['Word','Location']
        ResumeWords = ResumeWords.reset_index()
        ResumeWords = ResumeWords[['Word','Location']]
        df_ResumeWords=ResumeWords
        df_ResumeWords['Shift1'] =  df_ResumeWords['Word'].shift(-1)
        df_ResumeWords['Shift2'] =  df_ResumeWords['Word'].shift(-2) 
        df_ResumeWords['BiGram'] = df_ResumeWords['Word'] + ' ' + df_ResumeWords['Shift1']
        df_ResumeWords['TriGram'] = df_ResumeWords['Word'] + ' ' + df_ResumeWords['Shift1'] + ' ' + df_ResumeWords['Shift2']
        
        df_BiGrams = df_ResumeWords[['Location','BiGram']]
        df_TriGrams = df_ResumeWords[['Location','TriGram']]
        df_BiGrams['BiGram'] = df_BiGrams['BiGram'].str.rstrip()
        df_TriGrams['TriGram'] = df_TriGrams['TriGram'].str.rstrip()
        df_BiGrams = pd.merge(df_BiGrams,df_MainCluster[['Skill']],left_on='BiGram',right_on='Skill')
        df_TriGrams = pd.merge(df_TriGrams,df_MainCluster[['Skill']],left_on='TriGram',right_on='Skill')   
        
        #Get the location and previuos location for bigrams so that the words in those locations can be replaces with bigram
        df_BiGramsLocs= df_BiGrams[['Location']]
        df_BiGramsLocs['Location-1'] = df_BiGramsLocs['Location']+1
        df_BiGramsLocs['ID'] = df_BiGramsLocs.index
        df_BiGramsLocs= pd.melt(df_BiGramsLocs,id_vars=['ID'])
        df_BiGramsLocs = df_BiGramsLocs[['value']]
        df_BiGramsLocs.columns = ['Location']
        
        #Get the location,prev location and previuos to prev location for trigrams so that the words in those locations can be replaces with trigram
        df_TriGramsLocs= df_TriGrams[['Location']]
        df_TriGramsLocs['Location-1'] = df_TriGramsLocs['Location']+1
        df_TriGramsLocs['Location-2'] = df_TriGramsLocs['Location']+2
        df_TriGramsLocs['ID'] = df_TriGramsLocs.index
        df_TriGramsLocs= pd.melt(df_TriGramsLocs,id_vars=['ID'])
        df_TriGramsLocs = df_TriGramsLocs[['value']]
        df_TriGramsLocs.columns = ['Location']
        #Remove bigrams that are covered as part of trigrams based on location
        df_BiGramsLocs = df_BiGramsLocs[-df_BiGramsLocs['Location'].isin(df_TriGramsLocs['Location'].tolist())]  
        #Remove unigrams that are considered in bigram and trigram
        df_ResumeWords = df_ResumeWords[-df_ResumeWords['Location'].isin(df_BiGramsLocs['Location'].tolist())]
        df_ResumeWords = df_ResumeWords[-df_ResumeWords['Location'].isin(df_TriGramsLocs['Location'].tolist())]
        #Get the grams for the resime
        df_UniGrams = df_ResumeWords[['Word','Location']].reset_index()
        df_UniGrams = df_UniGrams[['Word','Location']]
        df_UniGrams.columns = ['Skill','Location']
        df_Grams = df_UniGrams.append(df_BiGrams[['Skill','Location']])
        df_Grams = df_Grams.append(df_TriGrams[['Skill','Location']])
        df_Words = df_Grams[['Skill']]
        df_Words.columns = ['Word'] 
        df_Words = df_Words[df_Words['Word']!=''].reset_index()
        df_Words = df_Words[['Word']]
        #Merge with skillMapping table to get the mapped word
        df_Words = pd.merge(df_Words,df_SkillMap,on='Word')
        df_Words = df_Words[['MappedTo']]
        df_Words.columns = ['Word']
        
        
        #Get the common grams between resume and main cluster
        df_WordsinMainCluster = pd.merge(df_Words,df_MainCluster,left_on='Word',right_on='Skill',how='inner') 
        #get the occurances of word in resume
        df_SkillCount =  pd.DataFrame(df_WordsinMainCluster.groupby(['Word'])['Word'].count())
        df_SkillCount['Skill'] = df_SkillCount.index
        df_SkillCount = df_SkillCount[['Word','Skill']]
        df_SkillCount.columns = ['Count','Skill']   
       
        ###Find the skills from text - ends          
        ### Step1: Identify Tops skills
        ### Use Associated Skills from Skillnorm, and aggregate to find the Top Skills in a Text
        lstSkillWords = df_SkillCount['Skill'].tolist()
        #Join existing skills with associated skills
        df_SkillWord_NM = df_SkillNorm[df_SkillNorm['AssociatedSkill'].isin(lstSkillWords)]
        #Filter to keep only significant skills 
        df_SkillWord_NM = df_SkillWord_NM[df_SkillWord_NM['WT_TotalOcc_Skill']>0.05]    
        
        #Core Skills are now weighted by their count of mention
        ##Start of Analysis by Core Skills mentioned
        t6 = df_MainCluster[df_MainCluster['Skill'].isin(lstSkillWords)]    
        t6 = t6[t6['WT_TotalOccurances']>0.05]
        t7 = pd.merge(t6,df_SkillCount,on='Skill')
        t7= t7[t7['WT_TotalOccurances']>.01]
        t7['Score'] = t7['WT_TotalOccurances'] * t7['Count']* (1-t7['WT_Spread'])
        t8 =  t7[['Skill','Score']]
        t8.columns = ['Skill','Score']
        ##End of analysis of Core skills 
        
        #Merge table with Maincluster to find the data profile of the Associated Skill
        t3 = df_SkillWord_NM
        Top10 = t8.nlargest(10,'Score')
        t3 = t3[t3['Skill'].isin(Top10['Skill'].tolist())]    
        t4 = t3[['Skill','AssociatedSkill','CoOccurranceRatio','WT_TotalOcc_Skill','WT_Spread_Skill','WT_TotalOcc_AssSkill','WT_Spread_Assskill']]
        t5 = pd.merge(t4,df_SkillCount,left_on='AssociatedSkill',right_on='Skill')
        t5 = t5[(t5['WT_TotalOcc_AssSkill']>.01) & (t5['WT_Spread_Assskill']<0.8) & (t5['CoOccurranceRatio']>0.05)]
        t5['Score']=  t5['WT_TotalOcc_AssSkill']*t5['Count']*t5['CoOccurranceRatio']
        #Aggregate the score by sum and find top skills
        t6 = t5.groupby('Skill_x').agg({'Score':'sum'}).reset_index()  
        df_TopSkill = t6
        df_TopSkill.columns = ['Skill','Score']
        #End of analysis by Association            
        
        #Summarize results to derive top skills
        #Append the above table to Top skill table and group by Sum of Score
        df_TopSkill = df_TopSkill.append(t8)
        df_TopSkill = df_TopSkill.groupby('Skill').agg({'Score':'sum'}).reset_index()                 
        df_TopSkill = df_TopSkill[df_TopSkill['Skill']!='none'].reset_index()
        df_TopSkill_10 = df_TopSkill.nlargest(10,'Score')
    
        df_SkillIdentified = pd.merge(df_SkillCount,df_MainCluster,on='Skill')
        df_SkillIdentified = df_SkillIdentified[df_SkillIdentified['WT_TotalOccurances']>.001]
        df_SkillIdentified['Score'] = df_SkillIdentified['WT_TotalOccurances'] * df_SkillIdentified['Count']
        return df_SkillCount,df_SkillIdentified,df_TopSkill_10
    except:
        print("not able to get skills")
        f=open('SkillMissing.txt','a')
        f.writelines('\n' + CurrentFile)
        f.close()
        df_SkillIdentified = pd.DataFrame(columns=['Count', 'Skill', 'TotalOccurances', 'AssociatedSkillSpread', 'WT_TotalOccurances', 'WT_Spread', 'WT_Total', 'Score'])
        df_TopSkill_10 = pd.DataFrame(columns=['index', 'Skill', 'Score'])
        df_SkillCount = pd.DataFrame(columns=['Count', 'Skill'])
        return df_SkillCount,df_SkillIdentified,df_TopSkill_10
def ProfileSensing(CurrentFile,sResumeID,sDocText, phoneNum, OrganizationID):    
    try:
        mydb = mysql.connector.connect(**pephire_db)
        mycursor = mydb.cursor()
        #If the same phone number exists in DB, delete that profile from DB
        if phoneNum:
            df = pd.read_sql(f"SELECT * FROM pephire.candidates WHERE phone='{phoneNum}' and organization_id='{OrganizationID}';", mydb)
            if not df.empty:
                resume_id = df['id'].iloc[0]
                mycursor.execute(f"delete from pephire.candidates WHERE phone='{phoneNum}' and id='{resume_id}';")
                mydb.commit()
                mycursor.execute(f"delete from pephire.candidateskills WHERE id='{resume_id}';")
                mydb.commit()
                mycursor.execute(f"delete from pephire.candidate_companies WHERE candidate_id='{resume_id}';")
                mydb.commit()
                mycursor.execute(f"delete from pephire.candidatesubskills WHERE id='{resume_id}';")
                mydb.commit()
                mycursor.execute(f"delete from pephire.candidate__tags WHERE candidate_id='{resume_id}';")
                mydb.commit()
                #print('FOUND EXISTING ENTRY WITH SAME PHONE NUM. DELETD OLD ENTRY')
        df_SkillCount,df_TopSkills,df_Top10Skills = GetSkill_New(CurrentFile,sDocText)                 
        df_Top10Skills['SkillScore'] = df_Top10Skills.apply(lambda x: f"{x['Skill']}<>{round(x['Score'], 2)}", axis=1)
        df_TopSkills['SkillScore'] = df_TopSkills.apply(lambda x: f"{x['Skill']}<>{round(x['Score'], 2)}", axis=1)            
        skillsTxt = '|'.join(df_Top10Skills['SkillScore'].to_list())
        subskillsTxt = '|'.join(df_TopSkills['SkillScore'].to_list())            
        SkillsDf = pd.DataFrame(columns = ['Skills','SubSkills'])
        SkillsDf.loc[0] = [skillsTxt, subskillsTxt]
        mydb.close()
        return SkillsDf        
        #return 2
    except:
        try:
            mydb.close()
        except:
            pass
        return pd.DataFrame(columns = ['Skills','SubSkills'])
    
import pandas as pd
from config import pephire_db_stat
import mysql.connector

mydb_static = mysql.connector.connect(**pephire_db_stat)
RoleCategoryMapping =pd.read_sql("select * from pephire_static.rolealias", mydb_static)
mydb_static.close()

#Read the config table for roles and role categories 

RoleCategoryMapping = RoleCategoryMapping.drop_duplicates()
RoleCategoryMapping['Role'] = RoleCategoryMapping['Role'].str.lower()
RoleCategoryMapping['Role Category'] = RoleCategoryMapping['Role Category'].str.lower()
RoleCategoryMapping['Alias'] = RoleCategoryMapping['Alias'].str.lower()



RoleUniAlias = RoleCategoryMapping[['Alias']]
RoleUniAlias = RoleUniAlias.Alias.str.split(" ",expand=True,)
RoleUniAlias['Row'] = RoleUniAlias.index
RoleUniAlias = RoleUniAlias.melt(id_vars=['Row'], var_name='var', value_name='Alias')
RoleUniAlias = RoleUniAlias[['Alias']].drop_duplicates().reset_index()
RoleUniAlias = RoleUniAlias[['Alias']]
RoleUniAlias = RoleUniAlias.dropna()
RoleUniAlias = RoleUniAlias[RoleUniAlias['Alias'].str.len()>1]


### Mithun Attention!
#1. We have got the Bigrams in a table. We need to save the location of the Bigrams
#2. Merge the Resume words, with location with Alias as a whole. 
#3. We now need to remove duplication. Example: If we considered, Performance tester, No need to consider Tester
#4. We do this by ensuring none of the unigram locations, are present in Bi Gram Locations, and Loc +/- 1
#5. Filter by above consider to get Unigrams, and append
#6. Now you have a clean table of roles, and their location
#7. Sort by location, count the total rows, weight diff = 1/(total rows), Decrement by Weight Diff (1, 1-wt, 1-2wt, 1-3wt)
#8. Now use this column for score. (add the role from role mapping (by Alias), and group by role, and sum of this weighted field)
        
def getAllRoles(CurrentFile,sDocText):
    try:
        ResumeWords = pd.DataFrame(re.split(' ',sDocText.lower().replace('\n',' ').replace('.',' ').replace(',',' ').replace('/',' ').replace(':',' ').replace('-',' ').replace('?',' ').replace('_',' ')))
        ResumeWords.columns = ['Word']
        ResumeWords['Location'] = ResumeWords.index
        ResumeWords.columns =['Word','Location']
        ResumeWords = ResumeWords.reset_index()
        ResumeWords = ResumeWords[['Word','Location']]
        #1. We have got the Bigrams in a table. We need to save the location of the Bigrams
        #get the common words between alias in Rolemapping and word in resume
        df_CommonUnigrams = pd.merge(ResumeWords,RoleUniAlias,left_on='Word',right_on='Alias')
        #ProfileBiGrams[['First','Second']] = ProfileBiGrams.Gram.str.split(" ",expand=True,)   
        ProfileBiGrams,ProfileTriGrams = pd.DataFrame(),pd.DataFrame()
        #For each common word in resume and Alias find the previous word in resume and form a bigram role
        for i, row in df_CommonUnigrams.iterrows():
            index = df_CommonUnigrams['Location'].iloc[i]            
            sWord = pd.DataFrame.from_dict({'Word':[ResumeWords.iloc[index-1][0] + ' ' + ResumeWords.iloc[index][0]],'Location':[index]})
            ProfileBiGrams = ProfileBiGrams.append(sWord)# save location i 
            sWord = pd.DataFrame.from_dict({'Word':[ResumeWords.iloc[index-2][0] + ' ' + ResumeWords.iloc[index-1][0] + ' ' + ResumeWords.iloc[index][0]],'Location':[index]})
            ProfileTriGrams = ProfileTriGrams.append(sWord)# save location i 
        ProfileBiGrams['Gram'] = ProfileBiGrams['Word']
        ProfileTriGrams['Gram'] = ProfileTriGrams['Word'] 

        #2. Merge the Resume words, with location with Alias as a whole.
        df_TriGramAlias = pd.merge(ProfileTriGrams,RoleCategoryMapping[['Alias']],left_on='Gram',right_on='Alias',how='inner')
        df_BiGramAlias = pd.merge(ProfileBiGrams,RoleCategoryMapping[['Alias']],left_on='Gram',right_on='Alias',how='inner')
        df_UniGramAlias = pd.merge(ResumeWords,RoleCategoryMapping[['Alias']],left_on='Word',right_on='Alias',how='inner')

        #Remove unigrams that are second word of bigram
        df_UniGramAlias = df_UniGramAlias[-df_UniGramAlias['Location'].isin(df_BiGramAlias['Location'].tolist())]
        df_UniGramAlias = df_UniGramAlias[-df_UniGramAlias['Location'].isin(df_TriGramAlias['Location'].tolist())]
        #6. Now you have a clean table of roles, and their location
        df_AllAlias = df_UniGramAlias.append(df_BiGramAlias[['Word','Location','Alias']])
        df_AllAlias = df_AllAlias.append(df_TriGramAlias[['Word','Location','Alias']])
        #ORder by location in ascending and get Role and return head 3 roles

        df_AllAlias = df_AllAlias.sort_values(by=['Location'])
        df_AllAlias = df_AllAlias.head(3)
        df_Top3Roles = pd.merge(df_AllAlias,RoleCategoryMapping[['Alias','Role']],on='Alias',how='left')
        df_Top3Roles['RoleScore'] = df_Top3Roles.apply(lambda x: f"{x['Alias']}<>{round(x['Location'], 2)}", axis=1)
        return df_Top3Roles
    except:         
         f=open('RoleMissing.txt','a')
         f.writelines('\n' + CurrentFile)
         f.close()
         return pd.DataFrame(columns=['Word','Location','Alias','Role','RoleScore'])

